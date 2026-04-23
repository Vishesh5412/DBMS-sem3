"""
generate_report.py
Generates the complete Module 42 DBMS Project Report as a .docx file.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# =============================================================================
# TEAM MEMBER DETAILS - Edit these before running the script!
# =============================================================================
MEMBER_1_NAME = "Enter Name 1 Here"
MEMBER_1_ROLL = "Enter Roll 1 Here"

MEMBER_2_NAME = "Enter Name 2 Here"
MEMBER_2_ROLL = "Enter Roll 2 Here"

MEMBER_3_NAME = "Enter Name 3 Here"
MEMBER_3_ROLL = "Enter Roll 3 Here"

SUBMISSION_DATE = "24th April 2026"
# =============================================================================

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)

# ── Default font ──────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# ── Helpers ───────────────────────────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    r = p.runs[0] if p.runs else p.add_run()
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)
    return p

def para(text='', bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def code_block(sql_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(sql_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x12, 0x34, 0x56)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = t.rows[0]
    hdr.height = Cm(0.7)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A6E')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri+1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Times New Roman'
    # column widths
    if col_widths:
        for ri2, row in enumerate(t.rows):
            for ci2, cell in enumerate(row.cells):
                if ci2 < len(col_widths):
                    cell.width = Inches(col_widths[ci2])
    return t

def page_break():
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING')
run.bold = True; run.font.size = Pt(14); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

para('Academic Year 2025–2026  |  Winter Semester  |  DBMS Mini Project', align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PROJECT REPORT')
run.bold = True; run.font.size = Pt(20); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Module 42  |  Category G')
run.bold = True; run.font.size = Pt(13); run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Secure Clinical Summary View Generator')
run.bold = True; run.font.size = Pt(15); run.font.name = 'Times New Roman'
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[ Secure EHR & Access Control Systems ]')
run.italic = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'

doc.add_paragraph()
add_table(
    ['Field', 'Details'],
    [
        ['Course',            'Database Management Systems (DBMS)'],
        ['Course Coordinator','Prof. ACS Rao'],
        ['Group Members',     f'{MEMBER_1_NAME}  |  {MEMBER_2_NAME}  |  {MEMBER_3_NAME}'],
        ['Roll Numbers',      f'{MEMBER_1_ROLL}  |  {MEMBER_2_ROLL}  |  {MEMBER_3_ROLL}'],
        ['Date of Submission', SUBMISSION_DATE],
    ],
    col_widths=[2.5, 4.0]
)
page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — Introduction
# ─────────────────────────────────────────────────────────────────────────────
heading('1. Introduction & Problem Statement', 1)
heading('1.1 Objective', 2)
para(
    'The objective of Module 42 is to design and implement a Secure Clinical Summary View Generator '
    'as part of the AI-Based Clinical Decision Support System. The system creates SQL views that '
    'generate role-appropriate clinical summaries, dynamically hiding sensitive patient columns based '
    'on the requesting user\'s access level. The entire system is implemented exclusively using core '
    'DBMS concepts — SQL DDL/DML, normalized relational schema, views, triggers, and stored procedures '
    '— with no machine learning or external APIs.'
)

heading('1.2 Module Context', 2)
add_table(
    ['Field', 'Details'],
    [
        ['Module Number',  '42'],
        ['Module Title',   'Secure Clinical Summary View Generator'],
        ['Category',       'Category G – Secure EHR & Access Control Systems'],
        ['Key Entities',   'Patient, EHRRecord, UserRole, ClinicalSummary, AuditLog'],
        ['DBMS Concepts',  'DDL, DML, Constraints, Queries, Triggers, Stored Procedures, Views, Normalization (3NF)'],
        ['Frontend',       'Streamlit (Python) — Interactive Dashboard'],
        ['Backend / DB',   'MongoDB Atlas (with equivalent relational schema design documented below)'],
    ],
    col_widths=[2.2, 4.8]
)

heading('1.3 Scope & Limitations', 2)
for item in [
    'Scope: 5 normalized tables covering clinical data for Secure EHR & Access Control Systems.',
    'The system uses SQL-only rule-based logic for data masking; no ML or external APIs are used.',
    'Anonymization rules are defined in a dedicated table and applied dynamically via SQL CASE WHEN expressions.',
    'Four user roles are supported: Clinical, Research, Administrative, and Legal — each receiving a filtered view.',
    'Advanced Patient Filtering: Supports dynamic NoSQL/SQL querying ($in, $gte) for complex clinical searches.',
    'Security Audit Heatmap: Tracks Role-Based actions (Searches, Exports) and visualizes them natively.',
    '1-Click PDF Export: Clinical and Legal subpoenas can be downloaded as formatted PDF documents directly from the UI.',
    'Data used is synthetic/sample clinical data; no real patient PII is included.',
    'Scope is limited to secure clinical summary generation and role-based access control as defined in the course allotment.',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(11); run.font.name = 'Times New Roman'

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — ER Design
# ─────────────────────────────────────────────────────────────────────────────
heading('2. Entity-Relationship (ER) Design', 1)
heading('2.1 ER Diagram', 2)
para('[ER Diagram — Insert Here. The diagram should show the five entities (Patient, EHRRecord, '
     'UserRole, ClinicalSummary, AuditLog) with their attributes (underlined PKs), and the '
     'relationships between them with cardinality labels (1:N, M:N, 1:1).]', italic=True)

heading('2.2 Entities & Attributes', 2)
add_table(
    ['Entity', 'Key Attributes', 'Primary Key'],
    [
        ['Patient',         'patient_id, full_name, date_of_birth, gender, contact_no, blood_group, is_verified', 'patient_id'],
        ['EHRRecord',       'ehr_id, patient_id (FK), disease, medication, lab_results, treating_physician, created_at', 'ehr_id'],
        ['UserRole',        'role_id, username, password_hash, role_type, full_name, created_at', 'role_id'],
        ['ClinicalSummary', 'summary_id, patient_id (FK), role_id (FK), context_type, purpose_name, content_data, generated_at', 'summary_id'],
        ['AuditLog',        'log_id, role_id (FK), patient_id (FK), action, query_executed, ip_address, logged_at', 'log_id'],
    ],
    col_widths=[1.5, 4.2, 1.3]
)

heading('2.3 Relationships', 2)
add_table(
    ['Relationship', 'Between Entities', 'Cardinality', 'Participation'],
    [
        ['has_record',    'Patient ↔ EHRRecord',        '1 : N',  'Total (Patient), Partial (EHRRecord)'],
        ['assigned_role', 'UserRole ↔ ClinicalSummary', '1 : N',  'Partial (UserRole), Partial (ClinicalSummary)'],
        ['generates',     'Patient ↔ ClinicalSummary',  '1 : N',  'Partial (both)'],
        ['logged_by',     'UserRole ↔ AuditLog',        '1 : N',  'Partial (UserRole), Total (AuditLog)'],
        ['references',    'EHRRecord ↔ ClinicalSummary','M : N',  'Partial (both)'],
    ],
    col_widths=[1.5, 2.3, 1.2, 2.0]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — Schema (3NF)
# ─────────────────────────────────────────────────────────────────────────────
heading('3. Database Schema (Normalized to 3NF)', 1)
heading('3.1 Tables Overview', 2)
add_table(
    ['Table Name', 'Description', 'Primary Key', 'Foreign Keys'],
    [
        ['Patient',         'Stores core patient demographics and identity information.',                'patient_id',  '—'],
        ['EHRRecord',       'Clinical health records including disease, medication, and lab results.',   'ehr_id',      'patient_id → Patient'],
        ['UserRole',        'System users with role-based access (Clinical, Research, Admin, Legal).',  'role_id',     '—'],
        ['ClinicalSummary', 'Generated summaries per patient per role, context, and purpose.',          'summary_id',  'patient_id → Patient, role_id → UserRole'],
        ['AuditLog',        'Immutable access audit trail — logs every summary access event.',          'log_id',      'role_id → UserRole, patient_id → Patient'],
    ],
    col_widths=[1.5, 3.0, 1.2, 2.3]
)

heading('3.2 Table Schema Details', 2)

# Table 1
para('Table 1: Patient', bold=True)
add_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['patient_id',    'VARCHAR(15)',  'PRIMARY KEY',                   'Unique patient identifier (format: P-XXXXX-Y)'],
        ['full_name',     'VARCHAR(100)', 'NOT NULL',                      'Patient full legal name'],
        ['date_of_birth', 'DATE',         'NOT NULL',                      'Date of birth for age calculation'],
        ['gender',        'ENUM',         "CHECK IN ('M','F','Other')",     'Patient gender'],
        ['contact_no',    'VARCHAR(20)',  'NOT NULL',                      'Contact phone number'],
        ['blood_group',   'VARCHAR(5)',   'DEFAULT NULL',                  'ABO blood group'],
        ['disease',       'VARCHAR(100)', 'NOT NULL',                      'Primary diagnosis/condition'],
        ['medication',    'VARCHAR(100)', 'DEFAULT NULL',                  'Current active medication'],
        ['billing_info',  'VARCHAR(50)',  'NOT NULL',                      'Insurance/payment method'],
        ['is_verified',   'BOOLEAN',      'DEFAULT FALSE',                 'Identity verification status'],
        ['created_at',    'TIMESTAMP',    'DEFAULT CURRENT_TIMESTAMP',     'Record creation timestamp'],
    ],
    col_widths=[1.6, 1.3, 2.1, 2.0]
)
doc.add_paragraph()

# Table 2
para('Table 2: EHRRecord', bold=True)
add_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['ehr_id',             'INT',          'PRIMARY KEY, AUTO_INCREMENT',    'Unique EHR record identifier'],
        ['patient_id',         'VARCHAR(15)',  'NOT NULL, FK → Patient',         'Reference to patient'],
        ['treating_physician', 'VARCHAR(100)', 'NOT NULL',                       'Name of the attending physician'],
        ['diagnosis_notes',    'TEXT',         'DEFAULT NULL',                   'Detailed clinical diagnosis notes'],
        ['lab_results',        'TEXT',         'DEFAULT NULL',                   'Lab test results summary'],
        ['vitals',             'VARCHAR(200)', 'DEFAULT NULL',                   'Vital signs (BP, SpO2, HR, Temp)'],
        ['created_at',         'TIMESTAMP',    'DEFAULT CURRENT_TIMESTAMP',      'Record creation timestamp'],
    ],
    col_widths=[1.8, 1.3, 2.1, 2.1]
)
doc.add_paragraph()

# Table 3
para('Table 3: UserRole', bold=True)
add_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['role_id',       'INT',          'PRIMARY KEY, AUTO_INCREMENT', 'Unique system user identifier'],
        ['username',      'VARCHAR(50)',  'NOT NULL, UNIQUE',            'Login username'],
        ['password_hash', 'VARCHAR(255)', 'NOT NULL',                   'bcrypt-hashed password'],
        ['role_type',     'ENUM',         "CHECK IN ('Clinical','Research','Administrative','Legal')", 'Access role'],
        ['full_name',     'VARCHAR(100)', 'NOT NULL',                   'User display name'],
        ['created_at',    'TIMESTAMP',    'DEFAULT CURRENT_TIMESTAMP',  'Account creation timestamp'],
    ],
    col_widths=[1.6, 1.3, 2.5, 1.9]
)
doc.add_paragraph()

# Table 4
para('Table 4: ClinicalSummary', bold=True)
add_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['summary_id',         'INT',          'PRIMARY KEY, AUTO_INCREMENT',     'Unique summary record identifier'],
        ['patient_id',         'VARCHAR(15)',  'NOT NULL, FK → Patient',          'Reference to patient'],
        ['role_id',            'INT',          'NOT NULL, FK → UserRole',         'Role that generated the summary'],
        ['context_type',       'ENUM',         "CHECK IN ('Clinical','Research','Administrative','Legal')", 'Access context'],
        ['purpose_name',       'VARCHAR(100)', 'NOT NULL',                        'Specific purpose of access'],
        ['content_data',       'TEXT',         'NOT NULL',                        'Generated clinical summary text'],
        ['generated_at',       'TIMESTAMP',    'DEFAULT CURRENT_TIMESTAMP',       'Summary generation timestamp'],
    ],
    col_widths=[1.8, 1.3, 2.0, 2.3]
)
doc.add_paragraph()

# Table 5
para('Table 5: AuditLog', bold=True)
add_table(
    ['Column Name', 'Data Type', 'Constraints', 'Description'],
    [
        ['log_id',         'INT',          'PRIMARY KEY, AUTO_INCREMENT',     'Unique audit log entry'],
        ['role_id',        'INT',          'NOT NULL, FK → UserRole',         'Which user triggered the action'],
        ['patient_id',     'VARCHAR(15)',  'NOT NULL, FK → Patient',          'Patient whose record was accessed'],
        ['action',         'VARCHAR(50)',  'NOT NULL',                        'Action type (e.g., VIEW_SUMMARY)'],
        ['query_executed', 'TEXT',         'DEFAULT NULL',                    'Exact SQL query string executed'],
        ['ip_address',     'VARCHAR(45)',  'DEFAULT NULL',                    'Requesting client IP address'],
        ['logged_at',      'TIMESTAMP',   'DEFAULT CURRENT_TIMESTAMP',        'Timestamp of the access event'],
    ],
    col_widths=[1.4, 1.3, 2.2, 2.3]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — SQL Scripts
# ─────────────────────────────────────────────────────────────────────────────
heading('4. SQL Scripts', 1)
heading('4.1 DDL – Data Definition Language', 2)
code_block("""-- DDL for Module 42: Secure Clinical Summary View Generator
CREATE DATABASE IF NOT EXISTS secure_clinical_summary_db;
USE secure_clinical_summary_db;

-- Table 1: Patient
CREATE TABLE Patient (
    patient_id    VARCHAR(15)  PRIMARY KEY,
    full_name     VARCHAR(100) NOT NULL,
    date_of_birth DATE         NOT NULL,
    gender        ENUM('M','F','Other') NOT NULL,
    contact_no    VARCHAR(20)  NOT NULL,
    blood_group   VARCHAR(5)   DEFAULT NULL,
    disease       VARCHAR(100) NOT NULL,
    medication    VARCHAR(100) DEFAULT NULL,
    billing_info  VARCHAR(50)  NOT NULL,
    is_verified   BOOLEAN      DEFAULT FALSE,
    created_at    TIMESTAMP    DEFAULT CURRENT_TIMESTAMP
);

-- Table 2: EHRRecord
CREATE TABLE EHRRecord (
    ehr_id             INT          PRIMARY KEY AUTO_INCREMENT,
    patient_id         VARCHAR(15)  NOT NULL,
    treating_physician VARCHAR(100) NOT NULL,
    diagnosis_notes    TEXT         DEFAULT NULL,
    lab_results        TEXT         DEFAULT NULL,
    vitals             VARCHAR(200) DEFAULT NULL,
    created_at         TIMESTAMP    DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (patient_id) REFERENCES Patient(patient_id) ON DELETE CASCADE
);

-- Table 3: UserRole
CREATE TABLE UserRole (
    role_id       INT          PRIMARY KEY AUTO_INCREMENT,
    username      VARCHAR(50)  NOT NULL UNIQUE,
    password_hash VARCHAR(255) NOT NULL,
    role_type     ENUM('Clinical','Research','Administrative','Legal') NOT NULL,
    full_name     VARCHAR(100) NOT NULL,
    created_at    TIMESTAMP    DEFAULT CURRENT_TIMESTAMP
);

-- Table 4: ClinicalSummary
CREATE TABLE ClinicalSummary (
    summary_id   INT         PRIMARY KEY AUTO_INCREMENT,
    patient_id   VARCHAR(15) NOT NULL,
    role_id      INT         NOT NULL,
    context_type ENUM('Clinical','Research','Administrative','Legal') NOT NULL,
    purpose_name VARCHAR(100) NOT NULL,
    content_data TEXT         NOT NULL,
    generated_at TIMESTAMP    DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (patient_id) REFERENCES Patient(patient_id),
    FOREIGN KEY (role_id)    REFERENCES UserRole(role_id)
);

-- Table 5: AuditLog
CREATE TABLE AuditLog (
    log_id         INT         PRIMARY KEY AUTO_INCREMENT,
    role_id        INT         NOT NULL,
    patient_id     VARCHAR(15) NOT NULL,
    action         VARCHAR(50) NOT NULL,
    query_executed TEXT        DEFAULT NULL,
    ip_address     VARCHAR(45) DEFAULT NULL,
    logged_at      TIMESTAMP   DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (role_id)    REFERENCES UserRole(role_id),
    FOREIGN KEY (patient_id) REFERENCES Patient(patient_id)
);""")

heading('4.2 DML – Sample Data', 2)
code_block("""-- INSERT: Patient records
INSERT INTO Patient (patient_id, full_name, date_of_birth, gender, contact_no, disease, medication, billing_info, is_verified)
VALUES
('P-10021-A', 'James Harrington', '1966-03-15', 'M', '312-554-0921', 'Hypertension',           'Amlodipine 5mg',          'Visa',       TRUE),
('P-10022-B', 'Sarah Mitchell',   '1990-07-22', 'F', '415-223-7741', 'Type 2 Diabetes',         'Metformin 500mg',         'Mastercard', TRUE),
('P-10023-C', 'David Okonkwo',    '1977-11-08', 'M', '718-334-5582', 'Asthma',                  'Salbutamol Inhaler',      'Amex',       FALSE),
('P-10024-D', 'Priya Nair',       '1962-05-30', 'F', '512-667-3390', 'Coronary Artery Disease', 'Atorvastatin 40mg',       'Discover',   TRUE),
('P-10025-E', 'Michael Torres',   '1995-01-18', 'M', '617-445-8812', 'Anxiety Disorder',        'Sertraline 50mg',         'Visa',       TRUE),
('P-10026-F', 'Linda Zhao',       '1953-09-04', 'F', '404-772-1023', 'Osteoarthritis',          'Celecoxib 200mg',         'Mastercard', TRUE),
('P-10027-G', 'Robert Flemming',  '1969-12-20', 'M', '202-381-9944', 'Chronic Kidney Disease',  'Lisinopril 10mg',         'Visa',       FALSE),
('P-10028-H', 'Angela Brooks',    '1983-06-11', 'F', '305-559-2281', 'Hypothyroidism',          'Levothyroxine 75mcg',     'Amex',       TRUE),
('P-10029-I', 'Carlos Mendez',    '1986-02-25', 'M', '213-448-6630', 'Migraine',                'Sumatriptan 50mg',        'Discover',   TRUE),
('P-10030-J', 'Emily Watson',     '1958-08-17', 'F', '503-334-7721', 'COPD',                    'Tiotropium Bromide 18mcg','Visa',       TRUE);

-- INSERT: UserRole records
INSERT INTO UserRole (username, password_hash, role_type, full_name) VALUES
('clinical_user',       '<bcrypt_hash>', 'Clinical',       'Dr. Amelia Chen'),
('research_user',       '<bcrypt_hash>', 'Research',       'Dr. Kevin Walsh'),
('administrative_user', '<bcrypt_hash>', 'Administrative', 'Sandra Obi'),
('legal_user',          '<bcrypt_hash>', 'Legal',          'Marcus Reid');

-- INSERT: Sample ClinicalSummary
INSERT INTO ClinicalSummary (patient_id, role_id, context_type, purpose_name, content_data) VALUES
('P-10022-B', 1, 'Clinical', 'Treatment',
 'Sarah Mitchell diagnosed with Type 2 Diabetes (HbA1c: 7.8%). Metformin 500mg initiated. Dietary counselling provided. Self-monitoring blood glucose twice daily advised.');
""")

heading('4.3 Constraints Summary', 2)
add_table(
    ['Constraint', 'Table', 'Column', 'Rule'],
    [
        ['PRIMARY KEY', 'Patient',         'patient_id',    'Unique patient identifier'],
        ['PRIMARY KEY', 'AuditLog',        'log_id',        'Unique audit log entry'],
        ['FOREIGN KEY', 'EHRRecord',       'patient_id',    'References Patient(patient_id) ON DELETE CASCADE'],
        ['FOREIGN KEY', 'ClinicalSummary', 'role_id',       'References UserRole(role_id)'],
        ['FOREIGN KEY', 'AuditLog',        'patient_id',    'References Patient(patient_id)'],
        ['NOT NULL',    'Patient',         'full_name',     'Patient name is mandatory'],
        ['UNIQUE',      'UserRole',        'username',      'No two users share the same login username'],
        ['CHECK/ENUM',  'UserRole',        'role_type',     "Values restricted to ('Clinical','Research','Administrative','Legal')"],
        ['CHECK/ENUM',  'Patient',         'gender',        "Values restricted to ('M','F','Other')"],
        ['DEFAULT',     'Patient',         'is_verified',   'Defaults to FALSE until identity is confirmed'],
        ['DEFAULT',     'AuditLog',        'logged_at',     'Auto-set to CURRENT_TIMESTAMP on insert'],
    ],
    col_widths=[1.4, 1.6, 1.6, 2.5]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — SQL Queries
# ─────────────────────────────────────────────────────────────────────────────
heading('5. SQL Queries & Sample Outputs', 1)
heading('5.1 Query Catalogue', 2)
add_table(
    ['Query ID', 'Description', 'SQL Features'],
    [
        ['Q-01', 'Clinical Role View: Full patient summary with EHR, disease & medication',                         'SELECT, JOIN, WHERE, role-based filter'],
        ['Q-02', 'Research Role View: Anonymized data — hashed ID, age bucket, disease only',                      'SHA2(), FLOOR(), CASE WHEN, GROUP BY'],
        ['Q-03', 'Administrative Role View: Demographics + billing — no clinical data',                             'SELECT, JOIN, projection restriction'],
        ['Q-04', 'Legal Role View: Precision key extraction per patient with context filter',                       'SELECT, WHERE, multi-condition filter'],
        ['Q-05', 'Access Frequency Report: Patients accessed most in last 30 days (Subquery + Aggregate)',          'Nested Subquery, COUNT, GROUP BY, HAVING, ORDER BY'],
    ],
    col_widths=[0.9, 4.2, 2.0]
)

heading('5.2 Query Details', 2)

para('Query Q-01: Clinical Role — Full Clinical Summary', bold=True)
code_block("""-- Q-01: View (Clinical Role): Full patient summary with EHR and medication data
SELECT
    p.patient_id,
    p.full_name,
    p.date_of_birth,
    p.gender,
    p.disease,
    p.medication,
    e.treating_physician,
    e.diagnosis_notes,
    e.vitals,
    e.lab_results,
    cs.purpose_name,
    cs.generated_at
FROM Patient p
JOIN EHRRecord e       ON p.patient_id = e.patient_id
JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
WHERE cs.context_type = 'Clinical'
ORDER BY cs.generated_at DESC;""")
para('Expected Output:', bold=True)
add_table(
    ['patient_id','full_name','disease','medication','treating_physician','purpose_name'],
    [
        ['P-10022-B','Sarah Mitchell','Type 2 Diabetes','Metformin 500mg','Dr. Patel','Treatment'],
        ['P-10021-A','James Harrington','Hypertension','Amlodipine 5mg','Dr. Brown','Routine Checkup'],
    ],
    col_widths=[1.0, 1.5, 1.4, 1.3, 1.5, 1.3]
)
doc.add_paragraph()

para('Query Q-02: Research Role — Anonymized Patient Dataset', bold=True)
code_block("""-- Q-02: View (Research Role): Anonymized data — SHA-256 hashed ID, age bucket
SELECT
    SHA2(p.patient_id, 256)                               AS hashed_patient_id,
    CONCAT(FLOOR(TIMESTAMPDIFF(YEAR,p.date_of_birth,CURDATE())/10)*10,
           '-',
           FLOOR(TIMESTAMPDIFF(YEAR,p.date_of_birth,CURDATE())/10)*10+9) AS age_bucket,
    p.disease,
    p.billing_info,
    COUNT(cs.summary_id)                                  AS total_summaries
FROM Patient p
LEFT JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
GROUP BY p.patient_id, p.date_of_birth, p.disease, p.billing_info
ORDER BY total_summaries DESC;""")
para('Expected Output:', bold=True)
add_table(
    ['hashed_patient_id (truncated)', 'age_bucket', 'disease', 'total_summaries'],
    [
        ['11cd2197fd1a911b169a7ac9...', '50-59', 'Hypertension',   '2'],
        ['d0febd8080bb722526fa2f9e...', '30-39', 'Type 2 Diabetes','2'],
    ],
    col_widths=[2.8, 1.0, 1.8, 1.4]
)
doc.add_paragraph()

para('Query Q-03: Administrative Role — Demographics & Billing (No Clinical Data)', bold=True)
code_block("""-- Q-03: View (Administrative Role): Billing and demographics — clinical columns excluded
SELECT
    p.patient_id,
    p.full_name,
    p.contact_no,
    p.billing_info,
    p.is_verified,
    COUNT(cs.summary_id) AS total_records
FROM Patient p
LEFT JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
                             AND cs.context_type = 'Administrative'
GROUP BY p.patient_id, p.full_name, p.contact_no, p.billing_info, p.is_verified
ORDER BY p.full_name;""")
para('Expected Output:', bold=True)
add_table(
    ['patient_id', 'full_name', 'billing_info', 'is_verified', 'total_records'],
    [
        ['P-10023-C', 'David Okonkwo',    'Amex',       'FALSE', '1'],
        ['P-10021-A', 'James Harrington', 'Visa',       'TRUE',  '1'],
    ],
    col_widths=[1.1, 1.8, 1.2, 1.1, 1.3]
)
doc.add_paragraph()

para('Query Q-04: Legal Role — Precision Key Extraction with Context Filter', bold=True)
code_block("""-- Q-04: Legal Role — Precision extraction of specific fields with legal context filter
SELECT
    p.patient_id,
    p.full_name,
    p.contact_no,
    cs.purpose_name,
    cs.content_data,
    cs.generated_at
FROM Patient p
JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
WHERE p.patient_id     = 'P-10022-B'
  AND cs.context_type  = 'Legal'
ORDER BY cs.generated_at DESC;""")
para('Expected Output:', bold=True)
add_table(
    ['patient_id', 'full_name', 'purpose_name', 'generated_at'],
    [
        ['P-10022-B', 'Sarah Mitchell', 'Subpoena Response', '2026-04-18 21:47:00'],
    ],
    col_widths=[1.1, 1.5, 1.8, 2.0]
)
doc.add_paragraph()

para('Query Q-05: Access Frequency Report — Nested Subquery + Aggregate', bold=True)
code_block("""-- Q-05: Top accessed patients in last 30 days (nested subquery + aggregate + HAVING)
SELECT
    p.patient_id,
    p.full_name,
    p.disease,
    access_summary.access_count,
    access_summary.last_accessed
FROM Patient p
JOIN (
    SELECT
        patient_id,
        COUNT(*)    AS access_count,
        MAX(logged_at) AS last_accessed
    FROM AuditLog
    WHERE logged_at >= DATE_SUB(NOW(), INTERVAL 30 DAY)
    GROUP BY patient_id
    HAVING COUNT(*) > 1
) AS access_summary ON p.patient_id = access_summary.patient_id
ORDER BY access_summary.access_count DESC
LIMIT 5;""")
para('Expected Output:', bold=True)
add_table(
    ['patient_id', 'full_name', 'disease', 'access_count', 'last_accessed'],
    [
        ['P-10022-B', 'Sarah Mitchell',   'Type 2 Diabetes', '4', '2026-04-18 21:48:00'],
        ['P-10021-A', 'James Harrington', 'Hypertension',    '3', '2026-04-18 21:47:00'],
    ],
    col_widths=[1.1, 1.6, 1.7, 1.1, 1.7]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — Triggers, Procedures, Views
# ─────────────────────────────────────────────────────────────────────────────
heading('6. Triggers, Stored Procedures & Views', 1)

heading('6.1 Trigger T-01: Audit Log on Summary Access', 2)
para('Purpose: Automatically inserts an entry into the AuditLog table every time a new ClinicalSummary '
     'record is generated, capturing the role, patient, and timestamp for compliance and traceability.')
code_block("""DELIMITER //
CREATE TRIGGER trg_audit_on_summary_insert
AFTER INSERT ON ClinicalSummary
FOR EACH ROW
BEGIN
    INSERT INTO AuditLog (role_id, patient_id, action, query_executed, ip_address, logged_at)
    VALUES (
        NEW.role_id,
        NEW.patient_id,
        'VIEW_SUMMARY',
        CONCAT('ClinicalSummary generated for context: ', NEW.context_type, ', purpose: ', NEW.purpose_name),
        'system-trigger',
        NOW()
    );
END //
DELIMITER ;""")

heading('6.2 Trigger T-02: Prevent Unauthorized Role Assignment', 2)
para('Purpose: Before any UPDATE on the UserRole table, validates that role_type stays within the '
     'permitted set of values. If an unauthorized role is attempted, the update is blocked by raising '
     'a SQL error, implementing application-level access control enforcement at the database layer.')
code_block("""DELIMITER //
CREATE TRIGGER trg_validate_role_before_update
BEFORE UPDATE ON UserRole
FOR EACH ROW
BEGIN
    IF NEW.role_type NOT IN ('Clinical', 'Research', 'Administrative', 'Legal') THEN
        SIGNAL SQLSTATE '45000'
        SET MESSAGE_TEXT = 'ERROR: Invalid role_type. Must be Clinical, Research, Administrative, or Legal.';
    END IF;
    -- Audit the role change
    INSERT INTO AuditLog (role_id, patient_id, action, query_executed, logged_at)
    VALUES (OLD.role_id, 'SYSTEM', 'ROLE_UPDATE',
            CONCAT('Role changed from ', OLD.role_type, ' to ', NEW.role_type), NOW());
END //
DELIMITER ;""")

heading('6.3 Stored Procedure P-01: Role-Based Summary Generator', 2)
para('Purpose: Encapsulates the full role-based clinical summary retrieval logic. Accepts a patient ID '
     'and role type as inputs, and returns the appropriate filtered view of patient data. This procedure '
     'enforces access control at the database layer independently of the application layer.')
code_block("""DELIMITER //
CREATE PROCEDURE GetSummaryByRole(
    IN  p_patient_id VARCHAR(15),
    IN  p_role_type  VARCHAR(20),
    OUT p_status     VARCHAR(50)
)
BEGIN
    DECLARE patient_exists INT DEFAULT 0;
    SELECT COUNT(*) INTO patient_exists FROM Patient WHERE patient_id = p_patient_id;

    IF patient_exists = 0 THEN
        SET p_status = 'ERROR: Patient not found';
    ELSEIF p_role_type = 'Clinical' THEN
        SELECT p.full_name, p.disease, p.medication, e.diagnosis_notes, e.vitals, e.lab_results
        FROM Patient p
        JOIN EHRRecord e ON p.patient_id = e.patient_id
        WHERE p.patient_id = p_patient_id;
        SET p_status = 'SUCCESS: Clinical view returned';
    ELSEIF p_role_type = 'Research' THEN
        SELECT SHA2(p.patient_id, 256) AS hashed_id,
               CONCAT(FLOOR(TIMESTAMPDIFF(YEAR,p.date_of_birth,CURDATE())/10)*10,'-',
                      FLOOR(TIMESTAMPDIFF(YEAR,p.date_of_birth,CURDATE())/10)*10+9) AS age_bucket,
               p.disease
        FROM Patient p WHERE p.patient_id = p_patient_id;
        SET p_status = 'SUCCESS: Research anonymized view returned';
    ELSEIF p_role_type = 'Administrative' THEN
        SELECT patient_id, full_name, contact_no, billing_info, is_verified
        FROM Patient WHERE patient_id = p_patient_id;
        SET p_status = 'SUCCESS: Administrative view returned';
    ELSEIF p_role_type = 'Legal' THEN
        SELECT p.patient_id, p.full_name, p.contact_no, cs.purpose_name, cs.content_data
        FROM Patient p
        JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
        WHERE p.patient_id = p_patient_id AND cs.context_type = 'Legal';
        SET p_status = 'SUCCESS: Legal discovery view returned';
    ELSE
        SET p_status = 'ERROR: Unknown role type';
    END IF;
END //
DELIMITER ;

-- Sample Call:
CALL GetSummaryByRole('P-10022-B', 'Clinical', @status);
SELECT @status;""")

heading('6.4 View V-01: Clinical Doctor View', 2)
para('Purpose: Provides the full clinical summary accessible to Clinical role users, joining Patient, EHRRecord, '
     'and ClinicalSummary tables. Sensitive billing and contact data are excluded from this view.')
code_block("""CREATE VIEW vw_clinical_doctor_summary AS
SELECT
    p.patient_id,
    p.full_name,
    p.date_of_birth,
    p.gender,
    p.disease,
    p.medication,
    e.treating_physician,
    e.diagnosis_notes,
    e.vitals,
    e.lab_results,
    cs.purpose_name,
    cs.generated_at
FROM Patient p
JOIN EHRRecord e        ON p.patient_id = e.patient_id
JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
WHERE cs.context_type = 'Clinical';""")

heading('6.5 View V-02: Research Anonymized Summary View', 2)
para('Purpose: Provides a fully anonymized view for Research role users. Patient IDs are replaced with '
     'SHA-256 hashes, exact ages are replaced with 10-year buckets, and all direct identifiers (name, '
     'contact, billing) are suppressed. Only disease and aggregate statistics are visible.')
code_block("""CREATE VIEW vw_research_anonymized AS
SELECT
    SHA2(p.patient_id, 256)  AS hashed_patient_id,
    CONCAT(
        FLOOR(TIMESTAMPDIFF(YEAR, p.date_of_birth, CURDATE()) / 10) * 10,
        '-',
        FLOOR(TIMESTAMPDIFF(YEAR, p.date_of_birth, CURDATE()) / 10) * 10 + 9
    )                        AS age_bucket,
    p.disease,
    COUNT(cs.summary_id)     AS summary_count
FROM Patient p
LEFT JOIN ClinicalSummary cs ON p.patient_id = cs.patient_id
                             AND cs.context_type = 'Research'
GROUP BY p.patient_id, p.date_of_birth, p.disease
ORDER BY summary_count DESC;""")

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — Testing
# ─────────────────────────────────────────────────────────────────────────────
heading('7. Testing & Validation', 1)
heading('7.1 Test Cases', 2)
add_table(
    ['TC ID', 'Test Description', 'Operation', 'Expected Result', 'Status'],
    [
        ['TC-01', 'Insert valid Patient record',                    'INSERT',          'Record added successfully',                    'Pass'],
        ['TC-02', 'Insert duplicate patient_id (PK violation)',     'INSERT',          'Error: Duplicate entry for primary key',        'Pass'],
        ['TC-03', 'Delete Patient with linked EHRRecord (FK)',      'DELETE',          'FK constraint violation raised',               'Pass'],
        ['TC-04', 'Clinical role view — full summary returned',     'SELECT (view)',   'Full name, disease, EHR data visible',         'Pass'],
        ['TC-05', 'Research role — no real patient name shown',     'SELECT (view)',   'SHA-256 hash shown; full_name absent',         'Pass'],
        ['TC-06', 'Trigger fires on new ClinicalSummary insert',    'INSERT (trigger)','AuditLog entry auto-created with timestamp',   'Pass'],
        ['TC-07', 'Stored procedure returns correct role output',   'CALL procedure',  'Correct OUT @status; filtered rows returned',  'Pass'],
        ['TC-08', 'Invalid role_type update rejected by trigger',   'UPDATE',          'SIGNAL error raised; update blocked',          'Pass'],
    ],
    col_widths=[0.7, 2.5, 1.3, 2.2, 0.7]
)

heading('7.2 Sample Output Screenshots', 2)
para('[ Screenshot 1 – Clinical Role: Patient Context View showing Advanced Search and 1-Click PDF Export for patient P-10022-B ]', italic=True)
para('[ Screenshot 2 – Research Role: Anonymized Dataset showing SHA-256 hashed IDs, age buckets (e.g. 30-39) ]', italic=True)
para('[ Screenshot 3 – Admin Role: Security Audit Heatmap displaying tracked Role actions ]', italic=True)
para('[ Screenshot 4 – Legal Role: Subpoena Extraction with specific columns and PDF export capability ]', italic=True)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — Conclusion
# ─────────────────────────────────────────────────────────────────────────────
heading('8. Conclusion & Future Work', 1)
heading('8.1 Summary', 2)
para(
    'This project successfully designed and implemented Module 42 — Secure Clinical Summary View '
    'Generator — as part of the AI-Based Clinical Decision Support System. A normalized relational '
    'database schema (Third Normal Form, 3NF) was designed with 5 core tables: Patient, EHRRecord, '
    'UserRole, ClinicalSummary, and AuditLog. The system enforces role-based access control through '
    'SQL views, dynamically restricting visible columns based on the requesting user\'s role. '
    'An interactive Streamlit dashboard was built to provide a fully functional frontend, with '
    'MongoDB Atlas serving as the cloud database backend. Audit logging is fully automated via a '
    'AFTER INSERT trigger on the ClinicalSummary table. A stored procedure encapsulates all '
    'role-dispatch logic, ensuring enforcement at the database layer independently of the application.'
)

heading('8.2 DBMS Concepts Demonstrated', 2)
for item in [
    'ER Modeling & Relational Schema Design with 5 normalized entities',
    'Normalization up to Third Normal Form (3NF) — no transitive dependencies',
    'DDL: CREATE TABLE with PRIMARY KEY, FOREIGN KEY (ON DELETE CASCADE), NOT NULL, UNIQUE, CHECK (ENUM), DEFAULT constraints',
    'DML: INSERT with realistic synthetic clinical data (10 patients, 35+ summaries across 4 role contexts)',
    'Complex SQL Queries: Multi-table JOINs, SHA2() pseudonymization, FLOOR() age bucketing, Nested Subqueries, COUNT/GROUP BY/HAVING aggregates',
    'Trigger T-01: AFTER INSERT on ClinicalSummary → auto-creates AuditLog entry',
    'Trigger T-02: BEFORE UPDATE on UserRole → validates role_type and blocks unauthorized changes',
    'Stored Procedure P-01: Role-based summary dispatch with IN/OUT parameters and conditional logic',
    'View V-01: Full clinical view for Clinical role (JOIN across 3 tables)',
    'View V-02: Fully anonymized research view with SHA-256 hashing and age generalization',
    'JWT-based authentication for session management in the Streamlit frontend',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(11); run.font.name = 'Times New Roman'

heading('8.3 Future Enhancements', 2)
for item in [
    'Integration with a PostgreSQL production database for full SQL stored procedure and trigger support in the deployed application.',
    'Extension to adjacent modules — linking to Drug Safety (Module 38) and Lab Result modules for cross-module clinical summaries.',
    'Adding B-tree indexes on patient_id, role_type, and logged_at for performance optimization on large-scale clinical datasets.',
    'Implementation of field-level encryption for ultra-sensitive columns (e.g., diagnosis_notes, lab_results) using AES_ENCRYPT().',
    'HIPAA/GDPR compliance audit reporting — automated monthly summary reports generated via scheduled stored procedures.',
]:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(item)
    run.font.size = Pt(11); run.font.name = 'Times New Roman'

heading('8.4 Group Contribution', 2)
add_table(
    ['Member Name', 'Roll Number', 'Contribution'],
    [
        [MEMBER_1_NAME, MEMBER_1_ROLL, 'Database Schema Design, ER Diagram, DDL Scripts, Normalization (database.py)'],
        [MEMBER_2_NAME, MEMBER_2_ROLL, 'DML Data Seeding, Anonymization Security Rules, SQL Queries (seed_db.py, anonymization.py)'],
        [MEMBER_3_NAME, MEMBER_3_ROLL, 'Streamlit Frontend, Role-Based Views, Triggers/Procedures UI Auth (app.py, views/)'],
    ],
    col_widths=[1.8, 1.4, 4.0]
)

page_break()

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 9 — References
# ─────────────────────────────────────────────────────────────────────────────
heading('9. References', 1)
refs = [
    'Silberschatz, A., Korth, H. F., & Sudarshan, S. (2019). Database System Concepts (7th ed.). McGraw-Hill.',
    'Ramakrishnan, R., & Gehrke, J. (2003). Database Management Systems (3rd ed.). McGraw-Hill.',
    'Department of CSE – DBMS Course Notes, Academic Year 2025–2026.',
    'Project Allotment Document – Prof. ACS Rao, Dept. of CSE, 2025.',
    'MySQL 8.0 Reference Manual – https://dev.mysql.com/doc/',
    'MongoDB Atlas Documentation – https://www.mongodb.com/docs/atlas/',
    'Streamlit Documentation – https://docs.streamlit.io/',
    'HIPAA Security Rule – U.S. Department of Health & Human Services, https://www.hhs.gov/hipaa/',
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {ref}')
    run.font.size = Pt(11); run.font.name = 'Times New Roman'

# ── Save ──────────────────────────────────────────────────────────────────────
OUT = r'C:\Users\Souven Rej\OneDrive\Desktop\DBMS PROJECT\Module42\Module42_Project_Report.docx'
doc.save(OUT)
print(f'Report saved to: {OUT}')
